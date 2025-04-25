import streamlit as st
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

# Función para aplicar formato
def aplicar_formato(run, bold=False, underline=False):
    run.bold = bold
    run.underline = underline
    run.font.name = 'Arial'
    run.font.size = Pt(11)

# Generar declaración
def generar_declaracion(doc, consecutivo, nombre, cedula, matricula, apto, torre,
                        Direccion, Nproyecto, municipio, constructora, Nit):
    titulo = ("MINISTERIO DE MINAS Y ENERGIA\n"
              "DECLARACION DE CUMPLIMIENTO DEL REGLAMENTO\n"
              "TECNICO DE INSTALACIONES ELECTRICAS")
    p_titulo = doc.add_paragraph()
    run_titulo = p_titulo.add_run(titulo)
    aplicar_formato(run_titulo, bold=True)
    p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()
    p_consecutivo = doc.add_paragraph()
    p_consecutivo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run_consec = p_consecutivo.add_run(f"No. {consecutivo}")
    aplicar_formato(run_consec, bold=True)

    doc.add_paragraph()
    p_texto = doc.add_paragraph()
    p_texto.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    aplicar_formato(p_texto.add_run("Yo "))
    aplicar_formato(p_texto.add_run(nombre), bold=True, underline=True)
    aplicar_formato(p_texto.add_run(", mayor de edad, identificado con la CC. No. "))
    aplicar_formato(p_texto.add_run(cedula), bold=True, underline=True)
    aplicar_formato(p_texto.add_run(" en mi condición de "))
    aplicar_formato(p_texto.add_run("INGENIERO ELECTRICISTA"), bold=True, underline=True)
    aplicar_formato(p_texto.add_run(" portador de la matrícula profesional No. "))
    aplicar_formato(p_texto.add_run(matricula), bold=True, underline=True)
    aplicar_formato(p_texto.add_run(", declaro bajo la gravedad del juramento, que la instalación "))
    aplicar_formato(p_texto.add_run("DE CONSTRUCCION DE LAS REDES ELÉCTRICAS DESDE EL ARMARIO DE MEDIDORES HASTA LAS SALIDAS DE USO FINAL EN EL APARTAMENTO "),
                    bold=True, underline=True)
    aplicar_formato(p_texto.add_run(apto), bold=True, underline=True)
    aplicar_formato(p_texto.add_run(" localizado en la "))
    aplicar_formato(p_texto.add_run(Direccion), bold=True, underline=True)
    aplicar_formato(p_texto.add_run(" "))
    aplicar_formato(p_texto.add_run(torre), bold=True, underline=True)
    aplicar_formato(p_texto.add_run(", APTO "))
    aplicar_formato(p_texto.add_run(apto), bold=True, underline=True)
    aplicar_formato(p_texto.add_run(" DE APARTAMENTOS "), bold=True, underline=True)
    aplicar_formato(p_texto.add_run(Nproyecto), bold=True, underline=True)
    aplicar_formato(p_texto.add_run(" "))
    aplicar_formato(p_texto.add_run(torre), bold=True, underline=True)
    aplicar_formato(p_texto.add_run(" del municipio de "))
    aplicar_formato(p_texto.add_run(municipio), bold=True, underline=True)
    aplicar_formato(p_texto.add_run(", de propiedad de la "))
    aplicar_formato(p_texto.add_run(constructora), bold=True, underline=True)
    aplicar_formato(p_texto.add_run(" NIT "))
    aplicar_formato(p_texto.add_run(Nit), bold=True, underline=True)
    aplicar_formato(p_texto.add_run(" cuya construcción estuvo a mi cargo, cumple con todos y cada uno de los requisitos que le aplican establecidos en el Reglamento Técnico de Instalaciones Eléctricas RETIE..."))

    aplicar_formato(p_texto.add_run("\n\n(1)(X)...\n"))
    aplicar_formato(p_texto.add_run("O"), bold=True)
    aplicar_formato(p_texto.add_run("\n(2)..."))

    doc.add_page_break()

# Interfaz Streamlit
st.title("Generador de Declaraciones RETIE")

with st.form("formulario"):
    st.subheader("Datos del Ingeniero y Proyecto")
    nombre = st.text_input("Nombre completo")
    cedula = st.text_input("Cédula con puntos")
    matricula = st.text_input("Matrícula profesional")
    torre = st.text_input("Nombre de la torre (Ej: Torre 1)")
    consecutivo_inicial = st.number_input("Consecutivo inicial", min_value=1, step=1)
    Nproyecto = st.text_input("Nombre del proyecto")
    municipio = st.text_input("Municipio")
    constructora = st.text_input("Constructora")
    Nit = st.text_input("NIT")
    Direccion = st.text_input("Dirección")

    st.subheader("Configuración de apartamentos")
    primer_piso_diferente = st.radio("¿El primer piso es diferente?", ["Sí", "No"])
    apartamentos = []

    if primer_piso_diferente == "Sí":
        cantidad_apartamentos_p1 = st.number_input("¿Cuántos apartamentos tiene el primer piso?", min_value=1, step=1)
        for i in range(cantidad_apartamentos_p1):
            apto = st.text_input(f"Apartamento del primer piso #{i+1}")
            if apto:
                apartamentos.append(apto)
        piso_inicio = 2
    else:
        piso_inicio = 1

    pisos_totales = st.number_input("¿Cuántos pisos tiene la torre?", min_value=piso_inicio, step=1)
    aptos_por_piso = st.number_input("¿Cuántos apartamentos por piso?", min_value=1, step=1)

    enviado = st.form_submit_button("Generar documento")

if enviado:
    for piso in range(piso_inicio, pisos_totales + 1):
        for num in range(1, aptos_por_piso + 1):
            apto_num = f"{piso}{num:02}"
            apartamentos.append(apto_num)

    doc = Document()
    consecutivo = consecutivo_inicial

    for apto in apartamentos:
        generar_declaracion(doc, f"{consecutivo:03}", nombre, cedula, matricula, apto, torre,
                            Direccion, Nproyecto, municipio, constructora, Nit)
        consecutivo += 1

    archivo_salida = "declaraciones_generadas.docx"
    doc.save(archivo_salida)
    with open(archivo_salida, "rb") as file:
        st.success(f"✅ Se generaron {len(apartamentos)} declaraciones.")
        st.download_button("Descargar documento", file, archivo_salida)
